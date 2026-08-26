from __future__ import annotations

import argparse
from pathlib import Path
from neoag.candidate_identity import identity_value
from .common import count_by, ensure_dir, markdown_table, read_tsv, strip_html_text
from .appm_review import parse_evidence_html


def load_top_candidates(path: str | None, n: int = 20) -> tuple[list[dict[str, str]], dict[str, int], dict[str, int]]:
    if not path or not Path(path).exists():
        return [], {}, {}
    _, rows = read_tsv(path)
    p_counts = dict(count_by(rows, "final_priority"))
    event_counts = dict(count_by(rows, "event_type"))
    top = []
    seen_events: set[str] = set()
    seen_peptide_hla: set[str] = set()
    for r in rows:
        event_id = r.get("event_id") or r.get("peptide_id") or ""
        peptide_hla_id = identity_value(r, "peptide_hla_id")
        if event_id in seen_events or peptide_hla_id in seen_peptide_hla:
            continue
        seen_events.add(event_id)
        seen_peptide_hla.add(peptide_hla_id)
        top.append({
            "grade": r.get("pipeline_r_grade") or r.get("evidence_grade") or r.get("final_priority", ""),
            "gene": r.get("gene", ""), "peptide": r.get("peptide", ""),
            "hla": r.get("hla_allele", ""), "type": r.get("event_type", ""),
            "evidence": r.get("evidence_rank_key") or r.get("rna_support_status", ""),
            "limitation": r.get("hard_failure_codes") or r.get("priority_cap_reason_codes") or "待实验验证",
            "use": (r.get("recommended_validation") or r.get("recommended_use", ""))[:160],
        })
        if len(top) >= n:
            break
    return top, p_counts, event_counts


def write_docx_if_available(path: Path, title: str, sections: list[tuple[str, str]]) -> bool:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception:
        return False
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph("研究性计算筛选结果；不构成临床诊断、治疗建议或疗效承诺。")
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        lines = [line for line in body.split("\n") if line.strip()]
        table_lines = [line for line in lines if line.strip().startswith("|")]
        if len(table_lines) >= 2:
            headers = [cell.strip() for cell in table_lines[0].strip("|").split("|")]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
            for line in table_lines[2:]:
                values = [cell.strip() for cell in line.strip("|").split("|")]
                cells = table.add_row().cells
                for index, value in enumerate(values[:len(headers)]):
                    cells[index].text = value
            for line in lines:
                if line not in table_lines:
                    doc.add_paragraph(line.strip())
        else:
            for line in lines:
                doc.add_paragraph(line.strip())
    doc.save(path)
    return True


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Generate patient-facing report from Project B evidence outputs")
    ap.add_argument("--recommendation")
    ap.add_argument("--netmhcpan42")
    ap.add_argument("--evidence-report")
    ap.add_argument("--ranking-compare-report")
    ap.add_argument("--appm-review")
    ap.add_argument("--ccf-review")
    ap.add_argument("--outdir", required=True)
    ap.add_argument("--title", default="肿瘤新抗原预测分析及检测结果报告｜患者沟通版")
    args = ap.parse_args(argv)
    outdir = ensure_dir(args.outdir)
    html_summary = parse_evidence_html(args.evidence_report)
    top, priority_counts, event_counts = load_top_candidates(args.recommendation, 20)
    n_rows = sum(priority_counts.values())
    unique_summary = ""
    if args.recommendation and Path(args.recommendation).exists():
        _, rows = read_tsv(args.recommendation)
        unique_peptides = len({r.get("peptide", "") for r in rows})
        unique_events = len({r.get("event_id", "") for r in rows})
        unique_summary = f"本次候选 peptide-HLA 组合共 {len(rows):,} 条，独特肽段约 {unique_peptides:,} 条，候选事件约 {unique_events:,} 个。"

    priority_text = "；".join([f"{k}={v}" for k, v in sorted(priority_counts.items())]) if priority_counts else "未提供 final_priority 统计"
    appm_bits = []
    for k in ["MHC-I status", "MHC-II status", "IFNG response status", "appm_call_confidence", "evidence_completeness", "immune_escape_risk"]:
        if k in html_summary:
            appm_bits.append(f"{k}: {html_summary[k]}")
    appm_text = "；".join(appm_bits) if appm_bits else "未提供完整 APPM evidence report。"
    top_md = markdown_table(top, ["grade", "gene", "type", "peptide", "hla", "evidence", "limitation", "use"], max_rows=10)

    sections = []
    sections.append(("阅读提示与重要说明", "本报告是基于测序数据和计算模型形成的研究性分析。计算候选不等于体内真实呈递、不等于T细胞能够识别，也不等于确定治疗方案。缺失证据按UNASSESSED/PARTIAL处理，不能解释为阴性。"))
    sections.append(("1. 报告摘要", f"{unique_summary}\n综合推荐分级为：{priority_text}。R1–R4或旧版A–D均表示研究验证优先级，不是临床疗效等级。"))
    sections.append(("2. 患者样本与测序数据", "样本配对、肿瘤/正常DNA深度、肿瘤纯度/倍性、RNA质量和参考版本应从run manifest读取；未提供项目保持未评估。"))
    sections.append(("3. HLA分型与抗原呈递条件", f"{appm_text}\nAPPM、HLA LOH和IFNG/JAK-STAT用于解释呈递条件，不能单独判断免疫治疗敏感、耐药或患者获益。"))
    sections.append(("4. 重点变异事件", "应按SNV、InDel、Fusion、Splice和DNA SV分别统计总体情况，并在每个赛道内按event_id去重展示Top5；不同赛道不按单一数值直接比较。"))
    sections.append(("5. 候选肽段Top10（按事件去重）", top_md))
    sections.append(("6. Top候选解读与实验建议", "SNV采用MT/WT成对短肽；InDel/frameshift采用novel-tail长肽或minigene；Fusion先做RT-PCR/Sanger断点确认；Splice先做targeted RNA精确junction确认，再进入长肽/minigene和T细胞功能实验。"))
    sections.append(("7. 分析方法", "事件标准化 → 肽段构建 → 呈递预测 → RNA/CCF/HLA-LOH/APPM/安全性证据 → hard fail与priority cap → R1–R4 → 同赛道Pareto → 确定性tie-break → 事件去重。"))
    sections.append(("8. 局限性与总体结论", "预测结合不等于体内呈递或临床疗效。正常表达、HSPC、正常蛋白组、ligandome和正常junction不完整时，安全性不得写为完全通过。"))
    sections.append(("附录A：R1–R4证据分层", "R1：关键证据较完整，第一批实验优先。\nR2：值得推进，但有一项或少量谨慎因素。\nR3：优先补RNA、事件真实性、安全性或呈递证据。\nR4：硬失败、明确风险或证据明显不足，当前暂不推进。"))
    sections.append(("附录B：术语说明", "HLA：细胞表面的肽段展示架。\nMT/WT：突变肽与正常肽的成对比较。\nCCF：估计携带事件的肿瘤细胞比例。\nAPPM：抗原加工与呈递机制。\njunction reads：精确跨越融合或异常剪接连接点的reads。"))
    sections.append(("附录C：附件与可追溯文件", "run_manifest.json；all_tool_results.tsv；ranked_events.evidence_consensus.tsv；ranked_peptides.evidence_consensus.tsv；ranked_peptides.weighted_baseline.tsv；validation_plan.tsv；evidence_conflicts.tsv。患者版不展示服务器绝对路径。"))

    md_lines = [f"# {args.title}", ""]
    for h, b in sections:
        md_lines += [f"## {h}", b, ""]
    (outdir / "patient_report.md").write_text("\n".join(md_lines), encoding="utf-8")
    html = "<html><head><meta charset='utf-8'><title>Patient Report</title></head><body>" + "\n".join(
        [f"<h1>{args.title}</h1>"] + [f"<h2>{h}</h2><pre style='white-space:pre-wrap'>{b}</pre>" for h, b in sections]
    ) + "</body></html>"
    (outdir / "patient_report.html").write_text(html, encoding="utf-8")
    wrote = write_docx_if_available(outdir / "patient_report.docx", args.title, sections)
    (outdir / "patient_report_outputs.txt").write_text("\n".join([str(outdir / "patient_report.md"), str(outdir / "patient_report.html"), str(outdir / "patient_report.docx") if wrote else "DOCX not generated: python-docx unavailable"]) + "\n", encoding="utf-8")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
